# -*- coding: utf-8 -*-
"""共有フォルダAI検索 — 共有フォルダ内の Excel/Word/PDF/テキストを読み取り、
Claude API で「どのファイルに何が書いてあるか」を日本語で回答する Flask アプリ。

設定は環境変数（または同じフォルダの .env）で行う:
  SEARCH_FOLDER      検索対象フォルダ（\\\\server\\share のようなUNCパス可）
  ANTHROPIC_API_KEY  Anthropic APIキー（未設定でもキーワード検索だけは動く）
  CACHE_FILE         インデックスキャッシュの保存先（省略時は .cache/index.json）
"""

import json
import os
import re
import threading
import time
from pathlib import Path

from flask import Flask, jsonify, render_template, request


def load_dotenv():
    """依存を増やさないための簡易 .env ローダー（KEY=VALUE 形式のみ）。"""
    env_path = Path(__file__).resolve().parent / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8-sig").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
SEARCH_FOLDER = os.environ.get("SEARCH_FOLDER", str(BASE_DIR / "sample-docs"))
CACHE_FILE = Path(os.environ.get("CACHE_FILE", str(BASE_DIR / ".cache" / "index.json")))
CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6")
PORT = int(os.environ.get("PORT", "5000"))
MAX_FILES_TO_AI = int(os.environ.get("MAX_FILES_TO_AI", "5"))
EXCERPT_CHARS = int(os.environ.get("EXCERPT_CHARS", "3000"))

app = Flask(__name__)


# ===== テキスト抽出 =====

def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    parts = []
    try:
        for ws in wb.worksheets:
            parts.append(f"【シート: {ws.title}】")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
    finally:
        wb.close()
    return "\n".join(parts)


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for page_no, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"【{page_no}ページ】\n{text}")
    return "\n".join(parts)


def extract_plain(path: Path) -> str:
    data = path.read_bytes()
    # 日本語Windowsの共有フォルダは cp932 のファイルが混在しがち
    for encoding in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


EXTRACTORS = {
    ".xlsx": extract_xlsx,
    ".xlsm": extract_xlsx,
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".txt": extract_plain,
    ".md": extract_plain,
    ".csv": extract_plain,
}


# ===== インデックス（ディスクキャッシュ付き。更新日時で差分更新） =====
#
# キーは SEARCH_FOLDER からの相対パス（絶対パスではない）。
# 共有フォルダやアプリ自体を別の場所へ移設しても、ファイルの更新日時さえ
# 保たれていれば（robocopy 等）キャッシュを再利用でき、フル再スキャンを避けられる。

index = {}  # 相対パス(posix) -> {"mtime": float, "text": str, "error": str|None}
index_lock = threading.Lock()
last_scan_at = None
last_scan_errors = []


def load_cache():
    if not CACHE_FILE.exists():
        return
    try:
        data = json.loads(CACHE_FILE.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return
    with index_lock:
        index.clear()
        index.update(data.get("index", {}))


def save_cache():
    """呼び出し元が index_lock を保持している前提。"""
    try:
        CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
        CACHE_FILE.write_text(json.dumps({"index": index}, ensure_ascii=False), encoding="utf-8")
    except OSError as e:
        last_scan_errors.append(f"キャッシュ保存失敗: {e}")


def scan_folder():
    global last_scan_at, last_scan_errors
    root = Path(SEARCH_FOLDER)
    if not root.exists():
        raise FileNotFoundError(f"検索対象フォルダが見つかりません: {SEARCH_FOLDER}")

    with index_lock:
        found = set()
        errors = []
        for path in root.rglob("*"):
            if not path.is_file():
                continue
            if path.suffix.lower() not in EXTRACTORS:
                continue
            if path.name.startswith("~$"):  # Officeの一時ファイル
                continue
            rel_key = path.relative_to(root).as_posix()
            found.add(rel_key)
            try:
                mtime = path.stat().st_mtime
            except OSError as e:
                errors.append(f"{path.name}: {e}")
                continue
            entry = index.get(rel_key)
            if entry and entry["mtime"] == mtime:
                continue
            try:
                text = EXTRACTORS[path.suffix.lower()](path)
                index[rel_key] = {"mtime": mtime, "text": text, "error": None}
            except Exception as e:  # 壊れたファイル等はスキップして続行
                index[rel_key] = {"mtime": mtime, "text": "", "error": str(e)}
                errors.append(f"{path.name}: {e}")
        for key in list(index):
            if key not in found:
                del index[key]
        last_scan_at = time.time()
        last_scan_errors = errors
        save_cache()
    return len(found), errors


# ===== キーワード検索（AIに渡す候補の絞り込み） =====

CJK_RUN = re.compile(r"[぀-ヿ㐀-鿿豈-﫿ｦ-ﾟ]+")
WORD = re.compile(r"[A-Za-z0-9_][A-Za-z0-9_\-]*")


def query_terms(query: str):
    """検索語を単語＋日本語バイグラムに分解する。"""
    terms = set()
    for token in WORD.findall(query):
        if len(token) >= 2:
            terms.add(token.lower())
    for run in CJK_RUN.findall(query):
        if len(run) == 1:
            terms.add(run)
        for i in range(len(run) - 1):
            terms.add(run[i : i + 2])
    return terms


def find_snippets(text: str, terms, limit=3, width=60):
    snippets = []
    lower = text.lower()
    used_ranges = []
    for term in sorted(terms, key=len, reverse=True):
        start = 0
        while len(snippets) < limit:
            pos = lower.find(term.lower(), start)
            if pos < 0:
                break
            start = pos + len(term)
            if any(a <= pos < b for a, b in used_ranges):
                continue
            s, e = max(0, pos - width), min(len(text), pos + len(term) + width)
            used_ranges.append((s, e))
            snippet = text[s:e].replace("\n", " ").strip()
            snippets.append(("…" if s > 0 else "") + snippet + ("…" if e < len(text) else ""))
        if len(snippets) >= limit:
            break
    return snippets


def rank_files(query: str):
    terms = query_terms(query)
    results = []
    with index_lock:
        for rel_path, entry in index.items():
            if entry["error"]:
                continue
            text = entry["text"]
            name = Path(rel_path).name
            lower = text.lower()
            name_lower = name.lower()
            score = 0
            for term in terms:
                t = term.lower()
                score += lower.count(t)
                if t in name_lower:
                    score += 5
            if score > 0:
                results.append({"path": rel_path, "score": score, "text": text})
    results.sort(key=lambda r: r["score"], reverse=True)
    return results, terms


def build_excerpt(text: str, terms, max_chars: int) -> str:
    """マッチ箇所を中心に、AIに渡す抜粋を組み立てる。"""
    if len(text) <= max_chars:
        return text
    lower = text.lower()
    positions = sorted(
        p for term in terms for p in [lower.find(term.lower())] if p >= 0
    )
    if not positions:
        return text[:max_chars]
    chunk = max_chars // max(1, min(len(positions), 3))
    parts, taken = [], 0
    for pos in positions[:3]:
        s = max(0, pos - chunk // 2)
        parts.append(text[s : s + chunk])
        taken += chunk
        if taken >= max_chars:
            break
    return "\n…\n".join(parts)


# ===== Claude API =====

def ask_claude(query: str, candidates):
    import anthropic

    sections = [f"=== ファイル: {c['path']} ===\n{c['excerpt']}" for c in candidates]
    docs = "\n\n".join(sections)

    prompt = f"""あなたは社内の共有フォルダを検索するアシスタントです。
以下は共有フォルダ内のファイルから抽出したテキストの抜粋です。ユーザーの質問に日本語で答えてください。

【回答のルール】
- どのファイルに書いてあるかを必ずファイル名付きで示す（例:「見積書2026.xlsx に記載があります」）
- 抜粋に書かれていないことは推測しない。見つからなければ「該当する記載は見つかりませんでした」と正直に言う
- 該当箇所の内容を簡潔に引用・要約する
- 複数のファイルに関連情報があれば、それぞれ挙げる

【ユーザーの質問】
{query}

【共有フォルダ内のファイル抜粋】
{docs}"""

    client = anthropic.Anthropic()
    try:
        message = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
    except anthropic.AuthenticationError:
        return None, "APIキーが無効です。.env の ANTHROPIC_API_KEY を確認してください。"
    except anthropic.RateLimitError:
        return None, "APIのレート制限に達しました。しばらく待ってから再実行してください。"
    except anthropic.APIStatusError as e:
        return None, f"APIエラー ({e.status_code}): {e.message}"
    except anthropic.APIConnectionError:
        return None, "Anthropic APIに接続できませんでした。ネットワークを確認してください。"

    answer = "".join(b.text for b in message.content if b.type == "text").strip()
    return answer, None


# ===== ルーティング =====

@app.get("/")
def home():
    return render_template("index.html")


@app.get("/api/status")
def api_status():
    with index_lock:
        file_count = len(index)
        error_count = sum(1 for e in index.values() if e["error"])
    return jsonify(
        {
            "folder": SEARCH_FOLDER,
            "folder_exists": Path(SEARCH_FOLDER).exists(),
            "file_count": file_count,
            "error_count": error_count,
            "last_scan_at": last_scan_at,
            "ai_ready": bool(os.environ.get("ANTHROPIC_API_KEY")),
            "model": CLAUDE_MODEL,
            "cache_file": str(CACHE_FILE),
        }
    )


@app.post("/api/rescan")
def api_rescan():
    try:
        count, errors = scan_folder()
    except FileNotFoundError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    return jsonify({"ok": True, "file_count": count, "errors": errors})


@app.post("/api/search")
def api_search():
    query = (request.get_json(silent=True) or {}).get("query", "").strip()
    if not query:
        return jsonify({"ok": False, "error": "質問を入力してください"}), 400
    if last_scan_at is None:
        try:
            scan_folder()
        except FileNotFoundError as e:
            return jsonify({"ok": False, "error": str(e)}), 400

    ranked, terms = rank_files(query)
    top = ranked[:MAX_FILES_TO_AI]

    files_payload = [
        {"path": r["path"], "score": r["score"], "snippets": find_snippets(r["text"], terms)}
        for r in top
    ]

    answer = None
    ai_error = None
    if not top:
        answer = "キーワードに一致するファイルが見つかりませんでした。別の言葉で試すか、「再スキャン」を実行してください。"
    elif os.environ.get("ANTHROPIC_API_KEY"):
        for r in top:
            r["excerpt"] = build_excerpt(r["text"], terms, EXCERPT_CHARS)
        answer, ai_error = ask_claude(query, top)
    else:
        ai_error = "ANTHROPIC_API_KEY が未設定のため、キーワード検索の結果のみ表示しています。"

    return jsonify({"ok": True, "answer": answer, "ai_error": ai_error, "files": files_payload})


if __name__ == "__main__":
    print(f"検索対象フォルダ: {SEARCH_FOLDER}")
    print(f"キャッシュファイル: {CACHE_FILE}")
    print(f"APIキー: {'✅ 設定済み' if os.environ.get('ANTHROPIC_API_KEY') else '⬜ 未設定（キーワード検索のみ）'}")
    load_cache()
    if index:
        print(f"キャッシュから {len(index)} ファイル分を読み込みました（差分のみ再スキャンします）")
    try:
        count, _ = scan_folder()
        print(f"スキャン完了: {count} ファイル")
    except FileNotFoundError as e:
        print(f"⚠️ {e}")
    app.run(host="127.0.0.1", port=PORT, debug=False)
